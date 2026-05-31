"""Convert markdown documents for production (preddiplomnaya) practice
into .docx per КубГУ ФТФ formatting (Times New Roman 14, A4 with 30/15/20/20 mm
margins, line spacing 1.5, first-line indent 1.25 cm).

Inputs (in source_md/):
  - individualnoe_zadanie.md  -> Individualnoe_zadanie_Achmiz.docx
  - dnevnik.md                -> Dnevnik_praktiki_Achmiz.docx
  - otchet.md                 -> Otchyot_Achmiz.docx
  - kharakteristika.md        -> Kharakteristika_Achmiz.docx

Outputs land in filled/.

Reuses formatting helpers conceptually similar to ../convert_to_docx.py but
does not produce a ВКР-style title page; the user inserts the title pages
manually from the provided templates.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
SRC_DIR = BASE_DIR / "source_md"
OUT_DIR = BASE_DIR / "filled"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_font(run, *, size=14, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def _add_field(paragraph, instr_text):
    run = paragraph.add_run()
    _set_font(run, size=14)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instr)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def _add_text_run(paragraph, text, *, size=14, bold=False, italic=False,
                  name="Times New Roman"):
    run = paragraph.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic, name=name)
    return run


def _parse_inline(paragraph, text, *, size=14):
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    for chunk in pattern.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            _add_text_run(paragraph, chunk[2:-2], size=size, bold=True)
        elif chunk.startswith("*") and chunk.endswith("*"):
            _add_text_run(paragraph, chunk[1:-1], size=size, italic=True)
        elif chunk.startswith("`") and chunk.endswith("`"):
            _add_text_run(paragraph, chunk[1:-1], size=size, name="Courier New")
        else:
            _add_text_run(paragraph, chunk, size=size)


# ---------------------------------------------------------------------------
# Document-level setup
# ---------------------------------------------------------------------------

def _setup_document(*, with_page_numbers=False):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    rFonts = style.element.rPr.rFonts
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Times New Roman")

    if with_page_numbers:
        _setup_page_numbers(section)
    _setup_heading_styles(doc)
    return doc


def _setup_page_numbers(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    _add_field(paragraph, "PAGE   \\* MERGEFORMAT")


def _setup_heading_styles(doc):
    for level, size in ((1, 16), (2, 15), (3, 14)):
        style = doc.styles[f"Heading {level}"]
        style.base_style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        rFonts = style.element.rPr.rFonts
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), "Times New Roman")
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
        pf.keep_with_next = True
        if level == 1:
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Cm(0)
        else:
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(1.25)


# ---------------------------------------------------------------------------
# Block elements
# ---------------------------------------------------------------------------

STRUCTURAL_TITLES = {
    "СОДЕРЖАНИЕ",
    "ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ",
    "ВВЕДЕНИЕ",
    "ЗАКЛЮЧЕНИЕ",
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
}


def _is_structural(text):
    return text.upper().strip() in STRUCTURAL_TITLES


def _add_centered_title(doc, text, *, size=16, bold=True, uppercase=False,
                        page_break=False, space_after=12):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.keep_with_next = True
    _add_text_run(p, text.upper() if uppercase else text, size=size, bold=bold)
    return p


def _add_section_heading(doc, text, level, *, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph(style=doc.styles[f"Heading {level}"])
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.keep_with_next = True
    size = 16 if level == 2 else 15
    _add_text_run(p, text, size=size, bold=True)
    return p


def _add_normal_paragraph(doc, text, *, indent=True):
    if not text.strip():
        return None
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25) if indent else Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _parse_inline(p, text)
    return p


def _add_bullet_list_item(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.left_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if text.startswith("— "):
        text = text[2:]
    _add_text_run(p, "— ", size=14)
    _parse_inline(p, text)
    return p


def _add_numbered_list_item(doc, num, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.left_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _add_text_run(p, f"{num} ", size=14)
    _parse_inline(p, text)
    return p


def _add_reference_item(doc, num, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.left_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _add_text_run(p, f"{num}. ", size=14)
    _parse_inline(p, text)
    return p


def _add_table_title(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.keep_with_next = True
    _parse_inline(p, text)
    return p


def _add_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    rows = [r + [""] * (col_count - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            text = cell_text.strip()
            for para in cell.paragraphs:
                if para.runs:
                    for run in list(para.runs):
                        run.text = ""
            para = cell.paragraphs[0]
            pf = para.paragraph_format
            pf.first_line_indent = Cm(0)
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.alignment = (WD_ALIGN_PARAGRAPH.CENTER if i == 0 or j == 0
                             else WD_ALIGN_PARAGRAPH.LEFT)
            _parse_inline(para, text, size=12)
            for run in para.runs:
                _set_font(run, size=12, bold=(i == 0))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Cm(0)
    spacer.paragraph_format.space_after = Pt(0)


def _parse_table_row(line):
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def _is_separator(row):
    return all(re.fullmatch(r":?-{2,}:?", c.strip()) for c in row if c.strip())


def _add_toc(doc):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    note = doc.add_paragraph()
    note.paragraph_format.first_line_indent = Cm(0)
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _add_text_run(
        note,
        "(после открытия выделите содержание и нажмите F9 — Word подставит "
        "фактические номера страниц)",
        size=11, italic=True,
    )


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

def convert(md_path: Path, docx_path: Path, *,
            with_page_numbers=False,
            section_page_breaks=False,
            first_h1_as_centered_title=True):
    """Convert one markdown into one .docx.

    Parameters
    ----------
    with_page_numbers : whether to stamp footer page numbers (used for the
        отчёт, not for short forms like the дневник).
    section_page_breaks : whether to start each top-level (## / # 1 ...) on
        a new page. Only true for the отчёт.
    first_h1_as_centered_title : the first '# ...' line becomes a centred
        boldface title at the start of the body (i.e. the document name).
    """
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = _setup_document(with_page_numbers=with_page_numbers)

    i = 0
    table_rows = []
    first_h1_seen = False
    in_references = False

    def flush_table():
        nonlocal table_rows
        if table_rows:
            _add_table(doc, table_rows)
            table_rows = []

    while i < len(lines):
        raw_line = lines[i]
        line = raw_line.rstrip()

        # ---- tables ----
        if line.strip().startswith("|"):
            row = _parse_table_row(line)
            if not _is_separator(row):
                table_rows.append(row)
            i += 1
            if i < len(lines) and lines[i].strip().startswith("|"):
                continue
            flush_table()
            continue
        flush_table()

        # ---- separators ----
        if line.strip() == "---":
            i += 1
            continue

        # ---- TOC marker ----
        if line.strip() == "[[TOC]]":
            _add_toc(doc)
            i += 1
            continue

        # ---- headings ----
        h_match = re.match(r"^(#{1,3})\s+(.*)", line)
        if h_match:
            level = len(h_match.group(1))
            heading = h_match.group(2).strip()
            heading = re.sub(r"\*\*(.+?)\*\*", r"\1", heading)
            heading = re.sub(r"\*(.+?)\*", r"\1", heading)

            if level == 1 and not first_h1_seen and first_h1_as_centered_title:
                first_h1_seen = True
                _add_centered_title(doc, heading, size=16, bold=True,
                                    uppercase=True, page_break=False,
                                    space_after=18)
                i += 1
                continue

            if level == 1:
                # Subsequent H1s become centred sub-titles.
                _add_centered_title(doc, heading, size=15, bold=True,
                                    uppercase=False, page_break=False,
                                    space_after=12)
                i += 1
                continue

            if _is_structural(heading):
                in_references = heading.upper() == "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"
                _add_centered_title(doc, heading, size=16, bold=True,
                                    uppercase=True,
                                    page_break=section_page_breaks,
                                    space_after=18)
                i += 1
                continue

            if level == 2:
                _add_section_heading(doc, heading, level=2,
                                     page_break=section_page_breaks)
            else:
                _add_section_heading(doc, heading, level=3, page_break=False)
            i += 1
            continue

        # ---- table title ----
        if re.match(r"^Таблица\s+\d", line.strip()):
            _add_table_title(doc, line.strip())
            i += 1
            continue

        # ---- empty ----
        if not line.strip():
            i += 1
            continue

        # ---- bullets ----
        if line.lstrip().startswith("— "):
            _add_bullet_list_item(doc, line.strip())
            i += 1
            continue
        bullet_match = re.match(r"^[-*]\s+(.*)", line)
        if bullet_match:
            _add_bullet_list_item(doc, "— " + bullet_match.group(1))
            i += 1
            continue

        # ---- numbered ----
        num_match = re.match(r"^(\d+)\.\s+(.*)", line)
        if num_match:
            num = int(num_match.group(1))
            body = num_match.group(2)
            if in_references:
                _add_reference_item(doc, num, body)
            else:
                _add_numbered_list_item(doc, f"{num}.", body)
            i += 1
            continue

        # ---- paragraph (consume continuation lines) ----
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            stripped = nxt.strip()
            if (not stripped
                    or stripped.startswith("#")
                    or stripped.startswith("|")
                    or stripped == "---"
                    or stripped == "[[TOC]]"
                    or re.match(r"^Таблица\s+\d", stripped)
                    or stripped.startswith("— ")
                    or re.match(r"^[-*]\s+", nxt)
                    or re.match(r"^\d+\.\s+", nxt)):
                break
            para_lines.append(nxt)
            i += 1
        text_block = " ".join(p.strip() for p in para_lines).strip()
        if text_block:
            _add_normal_paragraph(doc, text_block)
        continue

    flush_table()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Saved: {docx_path}")


JOBS = [
    {
        "src": "individualnoe_zadanie.md",
        "dst": "Individualnoe_zadanie_Achmiz.docx",
        "with_page_numbers": False,
        "section_page_breaks": False,
    },
    {
        "src": "dnevnik.md",
        "dst": "Dnevnik_praktiki_Achmiz.docx",
        "with_page_numbers": False,
        "section_page_breaks": False,
    },
    {
        "src": "otchet.md",
        "dst": "Otchyot_Achmiz.docx",
        "with_page_numbers": True,
        "section_page_breaks": True,
    },
    {
        "src": "kharakteristika.md",
        "dst": "Kharakteristika_Achmiz.docx",
        "with_page_numbers": False,
        "section_page_breaks": False,
    },
]


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    only = set(sys.argv[1:])
    for job in JOBS:
        if only and job["src"] not in only and job["dst"] not in only:
            continue
        src = SRC_DIR / job["src"]
        dst = OUT_DIR / job["dst"]
        convert(src, dst,
                with_page_numbers=job["with_page_numbers"],
                section_page_breaks=job["section_page_breaks"])


if __name__ == "__main__":
    main()
