"""Convert diploma.md to diploma.docx per КубГУ / ГОСТ 7.32 requirements.

Key formatting rules enforced (методические указания ФТФ КубГУ 2025):

* Page setup: A4, margins L=30 mm, R=15 mm, T=B=20 mm.
* Body text: Times New Roman 14 pt, line spacing 1.5, justify, first-line
  indent 1.25 cm.
* Structural headings (РЕФЕРАТ, СОДЕРЖАНИЕ, ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ,
  СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ, ПРИЛОЖЕНИЕ ...): uppercase, centred,
  bold, 16 pt, each on a new page, one blank line below.
* Section / sub-section headings (1, 1.1): first-line indent 1.25 cm,
  bold, 15-16 pt. Sections start on a new page.
* Figure captions ('Рисунок N — ...') and table titles ('Таблица N — ...'):
  figure caption centred without first-line indent, table title flush
  left without first-line indent.
* Page numbering: arabic, bottom-centre, starting from page 2 (title page
  is counted but not stamped).
* References use a hanging indent and 1.5 spacing.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
FIGURES_DIR = BASE_DIR / "figures"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_font(run, *, size=14, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    # Force the Cyrillic complex-script font as well — otherwise Word
    # silently falls back to Calibri for the Russian text.
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def _add_field(paragraph, instr_text):
    """Insert a Word field (PAGE, TOC, etc.) into a paragraph."""
    run = paragraph.add_run()
    _set_font(run, size=14)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instr)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def _add_text_run(paragraph, text, *, size=14, bold=False, italic=False,
                  name="Times New Roman"):
    run = paragraph.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic, name=name)
    return run


def _parse_inline(paragraph, text, *, size=14):
    """Render markdown bold/italic/code spans into runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    for chunk in pattern.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            _add_text_run(paragraph, chunk[2:-2], size=size, bold=True)
        elif chunk.startswith("*") and chunk.endswith("*"):
            _add_text_run(paragraph, chunk[1:-1], size=size, italic=True)
        elif chunk.startswith("`") and chunk.endswith("`"):
            _add_text_run(paragraph, chunk[1:-1], size=size, name="Courier New")
        else:
            _add_text_run(paragraph, chunk, size=size)


# ---------------------------------------------------------------------------
# Document-level setup
# ---------------------------------------------------------------------------

def _setup_document():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    # Title page must not show its page number.
    section.different_first_page_header_footer = True

    # Default style.
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    rFonts = style.element.rPr.rFonts
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Times New Roman")

    _setup_page_numbers(section)
    _setup_heading_styles(doc)
    return doc


def _setup_page_numbers(section):
    """Stamp arabic page numbers centred at the bottom, starting on page 2."""
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    _add_field(paragraph, "PAGE   \\* MERGEFORMAT")

    # First page footer is empty (different_first_page is True).
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    first_para = first_footer.paragraphs[0]
    first_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_para.paragraph_format.first_line_indent = Cm(0)
    # leave blank


def _setup_heading_styles(doc):
    """Make Heading 1/2/3 styles GOST-compliant so the TOC field works."""
    for level, size in ((1, 16), (2, 15), (3, 14)):
        style = doc.styles[f"Heading {level}"]
        style.base_style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        rFonts = style.element.rPr.rFonts
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), "Times New Roman")
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
        pf.keep_with_next = True
        if level == 1:
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Cm(0)
        else:
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(1.25)


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------

def _title_centered(doc, text, *, size=14, bold=False, space_after=0,
                    space_before=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if text:
        _add_text_run(p, text, size=size, bold=bold)
    return p


def _title_signature_line(doc, prefix, name):
    """Add a right-aligned line: 'Role ___________ I. O. Surname'."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _add_text_run(p, f"{prefix} ", size=14)
    _add_text_run(p, "_" * 28, size=14)
    _add_text_run(p, f" {name}", size=14)


def _add_title_page(doc):
    """Title page per Приложение Б of the methodology."""
    _title_centered(doc, "Министерство науки и высшего образования Российской Федерации")
    _title_centered(doc, "Федеральное государственное бюджетное образовательное учреждение")
    _title_centered(doc, "высшего образования")
    _title_centered(doc, "«Кубанский государственный университет»")
    _title_centered(doc, "(ФГБОУ ВО «КубГУ»)")
    _title_centered(doc, "")
    _title_centered(doc, "Факультет компьютерных технологий и прикладной математики")
    _title_centered(doc, "Кафедра информационных технологий")

    for _ in range(3):
        _title_centered(doc, "")

    # Допуск к защите
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _add_text_run(p, "Допустить к защите", size=14)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _add_text_run(p, "Заведующий кафедрой", size=14)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _add_text_run(p, "__________________ И. О. Фамилия", size=14)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _add_text_run(p, "«____» ______________ 2025 г.", size=14)

    for _ in range(2):
        _title_centered(doc, "")

    _title_centered(doc, "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", size=16, bold=True)
    _title_centered(doc, "(БАКАЛАВРСКАЯ РАБОТА)", size=16, bold=True, space_after=18)

    _title_centered(doc, "МОНИТОРИНГ И ТЕСТИРОВАНИЕ ВЕБ-СЕРВИСОВ",
                    size=16, bold=True, space_after=18)

    for _ in range(2):
        _title_centered(doc, "")

    # Author block (left-aligned label, right column).
    def _info_row(label, value):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        _add_text_run(p, f"{label} ", size=14)
        _add_text_run(p, "_" * 28, size=14)
        _add_text_run(p, f" {value}", size=14)

    _info_row("Работу выполнил", "Ачмиз Юрий Муратович")

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _add_text_run(p, "Направление подготовки  09.03.02  Информационные системы и технологии",
                  size=14)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _add_text_run(p, "Направленность (профиль)  Информационные системы и технологии",
                  size=14)

    for _ in range(2):
        _title_centered(doc, "")

    _info_row("Научный руководитель", "")
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _add_text_run(p, "канд. техн. наук, доцент ", size=14)
    _add_text_run(p, "_" * 22, size=14)
    _add_text_run(p, "  И. А. Парфёнова", size=14)

    for _ in range(1):
        _title_centered(doc, "")

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _add_text_run(p, "Нормоконтролёр ", size=14)
    _add_text_run(p, "_" * 30, size=14)
    _add_text_run(p, "  И. О. Фамилия", size=14)

    # Push to bottom.
    for _ in range(3):
        _title_centered(doc, "")

    _title_centered(doc, "Краснодар  2025", size=14)
    # The next structural heading (РЕФЕРАТ) will insert its own page break.


# ---------------------------------------------------------------------------
# Structural-element handling (РЕФЕРАТ, СОДЕРЖАНИЕ, ВВЕДЕНИЕ ...)
# ---------------------------------------------------------------------------

STRUCTURAL_TITLES = {
    "РЕФЕРАТ",
    "СОДЕРЖАНИЕ",
    "ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ",
    "ВВЕДЕНИЕ",
    "ЗАКЛЮЧЕНИЕ",
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
}


def _is_structural(text):
    upper = text.upper().strip()
    if upper in STRUCTURAL_TITLES:
        return True
    if upper.startswith("ПРИЛОЖЕНИЕ "):
        return True
    return False


def _add_structural_heading(doc, text, *, start_on_new_page=True,
                            style_name=None):
    if start_on_new_page:
        doc.add_page_break()
    p = doc.add_paragraph()
    if style_name:
        p.style = doc.styles[style_name]
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(18)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.keep_with_next = True
    _add_text_run(p, text.upper(), size=16, bold=True)
    return p


def _add_section_heading(doc, text, level):
    """Sections (1) start on a new page; sub-sections (1.1 / 1.1.1) don't."""
    if level == 2:
        doc.add_page_break()
    p = doc.add_paragraph(style=doc.styles[f"Heading {level}"])
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.keep_with_next = True
    size = 16 if level == 2 else 15
    _add_text_run(p, text, size=size, bold=True)
    return p


# ---------------------------------------------------------------------------
# Body elements
# ---------------------------------------------------------------------------

def _add_normal_paragraph(doc, text):
    if not text.strip():
        return None
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _parse_inline(p, text)
    return p


def _add_figure_caption(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _parse_inline(p, text)
    return p


def _add_table_title(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.keep_with_next = True
    _parse_inline(p, text)
    return p


def _add_figure(doc, image_path):
    if not image_path.exists():
        print(f"  WARN: figure not found: {image_path}")
        return None
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(14))
    return p


def _add_bullet_list_item(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.left_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if text.startswith("— "):
        text = text[2:]
    _add_text_run(p, "— ", size=14)
    _parse_inline(p, text)
    return p


def _add_numbered_list_item(doc, num, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.left_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _add_text_run(p, f"{num} ", size=14)
    _parse_inline(p, text)
    return p


def _add_reference_item(doc, num, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.left_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _add_text_run(p, f"{num}. ", size=14)
    _parse_inline(p, text)
    return p


def _add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(1.0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        _add_text_run(p, line or " ", size=10, name="Courier New")


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def _add_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    rows = [r + [""] * (col_count - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            text = cell_text.strip()
            for para in cell.paragraphs:
                if para.runs:
                    for run in list(para.runs):
                        run.text = ""
            para = cell.paragraphs[0]
            pf = para.paragraph_format
            pf.first_line_indent = Cm(0)
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.alignment = (WD_ALIGN_PARAGRAPH.CENTER if i == 0 or j == 0
                             else WD_ALIGN_PARAGRAPH.LEFT)
            _parse_inline(para, text, size=12)
            for run in para.runs:
                _set_font(run, size=12, bold=(i == 0))

    # Spacer after table.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = Cm(0)
    spacer.paragraph_format.space_after = Pt(0)


def _parse_table_row(line):
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def _is_separator(row):
    return all(re.fullmatch(r":?-{2,}:?", c.strip()) for c in row if c.strip())


# ---------------------------------------------------------------------------
# TOC
# ---------------------------------------------------------------------------

def _add_toc(doc):
    """Insert a TOC field that Word will populate on F9 / first open."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    # Note for the user that the TOC must be refreshed.
    note = doc.add_paragraph()
    note.paragraph_format.first_line_indent = Cm(0)
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _add_text_run(
        note,
        "(после открытия документа выделите содержание и нажмите F9 — "
        "Word подставит фактические номера страниц с отточием)",
        size=11, italic=True,
    )


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

FIGURE_RE = re.compile(r"^\[\[FIGURE:([^\]]+)\]\]\s*$")


def convert(md_path: str, docx_path: str):
    text = Path(md_path).read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = _setup_document()
    _add_title_page(doc)

    i = 0
    table_rows = []
    in_code = False
    code_lines = []
    in_references = False
    in_first_section = True  # the "# title" of the whole work — skipped

    def flush_table():
        nonlocal table_rows
        if table_rows:
            _add_table(doc, table_rows)
            table_rows = []

    while i < len(lines):
        raw_line = lines[i]
        line = raw_line.rstrip()

        # ---- code blocks ----
        if line.strip().startswith("```"):
            if in_code:
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw_line)
            i += 1
            continue

        # ---- tables ----
        if line.strip().startswith("|"):
            row = _parse_table_row(line)
            if not _is_separator(row):
                table_rows.append(row)
            i += 1
            if i < len(lines) and lines[i].strip().startswith("|"):
                continue
            flush_table()
            continue
        flush_table()

        # ---- page break marker ('---') is a no-op: headings already
        # insert their own page breaks where appropriate.
        if line.strip() == "---":
            i += 1
            continue

        # ---- TOC marker ----
        if line.strip() == "[[TOC]]":
            _add_toc(doc)
            i += 1
            continue

        # ---- figure marker ----
        fig_match = FIGURE_RE.match(line.strip())
        if fig_match:
            path = (BASE_DIR / fig_match.group(1)).resolve()
            _add_figure(doc, path)
            i += 1
            continue

        # ---- headings ----
        h_match = re.match(r"^(#{1,3})\s+(.*)", line)
        if h_match:
            level = len(h_match.group(1))
            heading = h_match.group(2).strip()
            heading = re.sub(r"\*\*(.+?)\*\*", r"\1", heading)
            heading = re.sub(r"\*(.+?)\*", r"\1", heading)

            if level == 1 and in_first_section:
                # Document title duplicated on page 1 — already on the title
                # page, skip it.
                in_first_section = False
                i += 1
                continue

            # Structural elements (РЕФЕРАТ, СОДЕРЖАНИЕ, ВВЕДЕНИЕ, ПРИЛОЖЕНИЕ ...)
            # are level 2 in the source markdown but must be rendered as
            # centred, uppercase, "structural" headings per ГОСТ 7.32.
            if _is_structural(heading):
                in_references = heading.upper() == "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"
                _add_structural_heading(doc, heading)
            elif level <= 2:
                _add_section_heading(doc, heading, level=2)
            else:
                _add_section_heading(doc, heading, level=3)
            i += 1
            continue

        # ---- figure caption ----
        if re.match(r"^Рисунок\s+\d", line.strip()):
            _add_figure_caption(doc, line.strip())
            i += 1
            continue

        # ---- table title ----
        if re.match(r"^Таблица\s+(\d|В\.)", line.strip()):
            _add_table_title(doc, line.strip())
            i += 1
            continue

        # ---- empty line ----
        if not line.strip():
            i += 1
            continue

        # ---- bullet list (em-dash) ----
        if line.lstrip().startswith("— "):
            _add_bullet_list_item(doc, line.strip())
            i += 1
            continue

        # ---- markdown bullet ----
        bullet_match = re.match(r"^[-*]\s+(.*)", line)
        if bullet_match:
            _add_bullet_list_item(doc, "— " + bullet_match.group(1))
            i += 1
            continue

        # ---- numbered list / references ----
        num_match = re.match(r"^(\d+)\.\s+(.*)", line)
        if num_match:
            num = int(num_match.group(1))
            body = num_match.group(2)
            if in_references:
                _add_reference_item(doc, num, body)
            else:
                _add_numbered_list_item(doc, f"{num}.", body)
            i += 1
            continue

        # ---- normal paragraph (collect continuation lines) ----
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            stripped = nxt.strip()
            if (not stripped
                    or stripped.startswith("#")
                    or stripped.startswith("|")
                    or stripped.startswith("```")
                    or stripped == "---"
                    or stripped == "[[TOC]]"
                    or FIGURE_RE.match(stripped)
                    or re.match(r"^Рисунок\s+\d", stripped)
                    or re.match(r"^Таблица\s+(\d|В\.)", stripped)
                    or stripped.startswith("— ")
                    or re.match(r"^[-*]\s+", nxt)
                    or re.match(r"^\d+\.\s+", nxt)):
                break
            para_lines.append(nxt)
            i += 1
        text_block = " ".join(p.strip() for p in para_lines).strip()
        if text_block:
            _add_normal_paragraph(doc, text_block)
        continue

    flush_table()
    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    convert(
        str(BASE_DIR / "diploma.md"),
        str(BASE_DIR / "diploma.docx"),
    )
